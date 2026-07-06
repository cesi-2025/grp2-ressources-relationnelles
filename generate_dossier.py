#!/usr/bin/env python3
"""
Générateur du Dossier Technique — (RE)Sources Relationnelles
CESI — Bloc 2 — Groupe 2 — INFCDAAL2
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime
import os

# =============================================================================
# CONFIGURATION
# =============================================================================

OUTPUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Dossier_Technique.docx")
FONT_BODY = "Calibri"
FONT_SIZE_BODY = Pt(11)
FONT_SIZE_H1 = Pt(18)
FONT_SIZE_H2 = Pt(14)
FONT_SIZE_H3 = Pt(12)
PRIMARY_COLOR = RGBColor(0x1D, 0x3A, 0x6D)
ACCENT_COLOR = RGBColor(0x4C, 0xAF, 0x50)
GRAY_COLOR = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# =============================================================================
# STYLES & FORMATTING HELPERS
# =============================================================================

def setup_styles():
    """Configure document styles."""
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = FONT_SIZE_BODY
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level, (size, bold, color) in enumerate([
        (FONT_SIZE_H1, True, PRIMARY_COLOR),
        (FONT_SIZE_H2, True, PRIMARY_COLOR),
        (FONT_SIZE_H3, True, RGBColor(0x33, 0x33, 0x33)),
    ], start=1):
        h = doc.styles[f'Heading {level}']
        h.font.name = FONT_BODY
        h.font.size = size
        h.font.bold = bold
        h.font.color.rgb = color
        h.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        h.paragraph_format.space_after = Pt(8)

def setup_margins():
    """Set page margins."""
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

def add_page_numbers():
    """Add page numbers in footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fld_char_begin)
        run2 = p.add_run()
        instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instr)
        run3 = p.add_run()
        fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fld_char_end)

def add_header_text():
    """Add header with project name."""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("(RE)Sources Relationnelles — Dossier Technique")
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY_COLOR
        run.font.italic = True

def para(text, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, style=None, font_size=None, color=None, space_after=None):
    """Add a paragraph with formatting."""
    p = doc.add_paragraph(style=style)
    p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size:
        run.font.size = font_size
    if color:
        run.font.color.rgb = color
    return p

def bullet(text, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_table(headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1D3A6D"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(9)
            # Alternate row shading
            if r_idx % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F6FA"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing
    return table

def visual_placeholder(vtype, section, content, tool, size, legend):
    """Add a visual placeholder box."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.text = ""

    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFF9E6"/>')
    cell._tc.get_or_add_tcPr().append(shading)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("[ESPACE RESERVE — A COMPLETER]\n\n")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0x77, 0x00)

    details = (
        f"Type      : {vtype}\n"
        f"Section   : {section}\n"
        f"Contenu   : {content}\n"
        f"Outil     : {tool}\n"
        f"Taille    : {size}\n"
        f"Legende   : {legend}"
    )
    run2 = p.add_run(details)
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run2.italic = True

    doc.add_paragraph()
    return tbl

def page_break():
    doc.add_page_break()


# =============================================================================
# DOCUMENT GENERATION
# =============================================================================

setup_styles()
setup_margins()

# ─────────────────────────────────────────────────────────────────────────────
# PAGE DE GARDE
# ─────────────────────────────────────────────────────────────────────────────

for _ in range(6):
    doc.add_paragraph()

para("(RE)Sources Relationnelles", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
     font_size=Pt(28), color=PRIMARY_COLOR)
doc.add_paragraph()
para("Dossier Technique", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
     font_size=Pt(18), color=GRAY_COLOR)
doc.add_paragraph()
para("Projet collaboratif — INFCDAAL2", alignment=WD_ALIGN_PARAGRAPH.CENTER,
     font_size=Pt(14), color=GRAY_COLOR)

for _ in range(3):
    doc.add_paragraph()

para("Plateforme de ressources pour l'amelioration\ndes relations entre citoyens",
     italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=Pt(12), color=GRAY_COLOR)

for _ in range(4):
    doc.add_paragraph()

# Team info
team_info = [
    ["Chamil MUSAYEV", "Backend — API, BDD, Auth, Securite, RGPD"],
    ["Jamal", "Frontend web — Next.js, Charte graphique, RGAA"],
    ["Romain CONDOMITTI", "Back-office — Admin, Auth frontend, Integration"],
    ["Leon HEU", "Mobile — React Native, Expo"],
]
add_table(
    ["Membre", "Perimetre"],
    team_info,
    col_widths=[6, 10]
)

para(f"CESI Ecole d'Ingenieurs — {datetime.date.today().strftime('%d/%m/%Y')}",
     alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=Pt(11), color=GRAY_COLOR)

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# TABLE DES MATIERES
# ─────────────────────────────────────────────────────────────────────────────

doc.add_heading("Table des matieres", level=1)
para("(Table des matieres a generer automatiquement dans Word : References > Table des matieres)",
     italic=True, color=GRAY_COLOR, font_size=Pt(10))

# Insert a TOC field
p_toc = doc.add_paragraph()
run_toc = p_toc.add_run()
fld_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
run_toc._r.append(fld_begin)
run_toc2 = p_toc.add_run()
instr_toc = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
run_toc2._r.append(instr_toc)
run_toc3 = p_toc.add_run()
fld_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
run_toc3._r.append(fld_end)

para("Astuce : Apres ouverture du document dans Word, faites clic droit sur la table ci-dessus > Mettre a jour les champs > Mettre a jour toute la table.",
     italic=True, color=GRAY_COLOR, font_size=Pt(9))

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 1. INTRODUCTION ET CONTEXTE
# ─────────────────────────────────────────────────────────────────────────────

doc.add_heading("1. Introduction et contexte du projet", level=1)

doc.add_heading("1.1 Contexte general", level=2)
para("Le projet (RE)Sources Relationnelles s'inscrit dans une demarche du Ministere des Solidarites et de la Sante visant a proposer une plateforme numerique de ressources pour ameliorer la qualite des relations interpersonnelles des citoyens. Inspire de la pyramide de Maslow, le projet repond aux besoins de securite, de sens et de developpement personnel a travers la qualite de nos liens relationnels : famille, couple, amis, collegues.")

para("Ce document technique presente l'ensemble des choix architecturaux, des solutions techniques retenues et de la demarche de developpement mise en oeuvre par notre equipe pour repondre au cahier des charges. Il couvre les deux applications developpees (web et mobile) ainsi que le plan de tests associe.")

doc.add_heading("1.2 Objectifs de la plateforme", level=2)
para("Les objectifs principaux identifies dans le CCTP sont :")
bullet("Mise a disposition de ressources de differents types pour les usagers (articles, videos, activites)")
bullet("Gestion dynamique du catalogue de ressources (categories, types de relations, types de ressources)")
bullet("Possibilite pour les citoyens de creer et partager leurs propres ressources")
bullet("Espace d'echange modere sur les ressources publiques (commentaires hierarchiques)")
bullet("Suivi de progression du parcours utilisateur (favoris, ressources exploitees, mises de cote)")
bullet("Espace statistiques pour les administrateurs (consultations, creations, exploitations)")

doc.add_heading("1.3 Acteurs du systeme", level=2)
para("Le systeme identifie cinq profils d'acteurs avec des niveaux de privileges distincts :")

add_table(
    ["Acteur", "Role", "Acces"],
    [
        ["Citoyen non connecte", "Consultation des ressources publiques", "Front-office (lecture seule)"],
        ["Citoyen connecte", "Creation de ressources, commentaires, favoris, progression", "Front-office (lecture/ecriture)"],
        ["Moderateur", "Validation des ressources et commentaires", "Front-office + moderation"],
        ["Administrateur", "Gestion du catalogue, statistiques, suspension", "Front-office + back-office"],
        ["Super-administrateur", "Creation de comptes privilegies", "Acces complet"],
    ],
    col_widths=[4, 6, 5]
)

doc.add_heading("1.4 Perimetre du document", level=2)
para("Ce dossier couvre les elements suivants, conformement aux consignes du projet collaboratif INFCDAAL2 :")
bullet("Analyse fonctionnelle et modelisation UML")
bullet("Choix technologiques argumentes pour chaque composant")
bullet("Architecture applicative detaillee (pattern MVC, API REST, base de donnees)")
bullet("Developpement backend : API Laravel, authentification, securite, RGPD")
bullet("Developpement frontend : Next.js, accessibilite RGAA, responsive design")
bullet("Application mobile : React Native / Expo")
bullet("Plan de tests : methodologie, cahier de tests, PV de recette")
bullet("Deploiement : Docker, CI/CD, environnement de production")
bullet("Organisation du projet : equipe, workflow Git, sprints")

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 2. ANALYSE FONCTIONNELLE
# ─────────────────────────────────────────────────────────────────────────────

doc.add_heading("2. Analyse fonctionnelle", level=1)

doc.add_heading("2.1 Diagramme des cas d'utilisation", level=2)
para("Le diagramme de cas d'utilisation ci-dessous represente les interactions entre les differents acteurs et le systeme. Il a ete elabore a partir des fonctionnalites listees dans le CCTP et reflete l'ensemble des besoins fonctionnels identifies.")

visual_placeholder(
    "Diagramme de cas d'utilisation UML",
    "2.1 Analyse fonctionnelle",
    "Diagramme montrant les 5 acteurs (Citoyen non connecte, Citoyen connecte, Moderateur, Administrateur, Super-admin) et leurs cas d'utilisation respectifs : consulter ressources, creer ressource, commenter, moderer, gerer catalogue, voir statistiques, creer comptes privilegies. Inclure les relations <<include>> et <<extend>>.",
    "draw.io / PlantUML / StarUML",
    "Pleine page",
    "Figure 1 — Diagramme de cas d'utilisation du systeme (RE)Sources Relationnelles"
)

para("Le diagramme fait apparaitre les principaux cas d'utilisation regroupes par domaine fonctionnel :")
bullet("Acces aux ressources : consultation publique, filtrage, tri, affichage du detail")
bullet("Gestion de compte : inscription, connexion, deconnexion, anonymisation (RGPD)")
bullet("Creation de contenu : creation/edition de ressources, soumission a la moderation")
bullet("Interactions : commentaires, reponses, ajout aux favoris, suivi de progression")
bullet("Administration : statistiques, suspension de ressources, gestion des categories")
bullet("Super-administration : creation de comptes moderateurs et administrateurs")

doc.add_heading("2.2 Analyse de la valeur", level=2)

visual_placeholder(
    "Diagramme bete a cornes",
    "2.2 Analyse de la valeur",
    "Diagramme bete a cornes repondant aux questions : A qui rend-il service ? (Citoyens, familles, couples) / Sur quoi agit-il ? (Relations interpersonnelles) / Dans quel but ? (Ameliorer la qualite des liens relationnels via des ressources adaptees)",
    "draw.io / PowerPoint",
    "Demi-page",
    "Figure 2 — Diagramme bete a cornes"
)

para("L'analyse de la valeur met en evidence que la plateforme rend service aux citoyens en agissant sur la qualite de leurs relations interpersonnelles. Le besoin fondamental est de fournir un acces simple, gratuit et modere a des ressources pedagogiques et pratiques.")

visual_placeholder(
    "Diagramme pieuvre (analyse fonctionnelle externe)",
    "2.2 Analyse de la valeur",
    "Diagramme pieuvre avec le produit au centre et les elements du milieu exterieur : Utilisateurs, Navigateur web, Smartphone, Base de donnees, Internet, Normes RGPD/RGAA, Ministere. Fonctions principales (FP) et fonctions contraintes (FC).",
    "draw.io / PowerPoint",
    "Demi-page",
    "Figure 3 — Diagramme pieuvre"
)

para("Les fonctions principales identifiees sont :")
bullet("FP1 : Permettre aux citoyens d'acceder a des ressources relationnelles via un navigateur web")
bullet("FP2 : Permettre aux citoyens de consulter et interagir avec les ressources depuis un smartphone")
bullet("FC1 : Respecter les normes d'accessibilite (RGAA) et de protection des donnees (RGPD)")
bullet("FC2 : S'adapter aux contraintes de performance et de disponibilite d'une plateforme publique")
bullet("FC3 : Etre compatible avec les navigateurs courants du marche (Chrome, Firefox, Safari, Edge)")

doc.add_heading("2.3 Liste des fonctionnalites implementees", level=2)
para("Le tableau suivant recapitule l'ensemble des fonctionnalites identifiees dans le CCTP et leur etat d'implementation dans la plateforme :")

add_table(
    ["Domaine fonctionnel", "Fonctionnalite", "Acteur", "Statut"],
    [
        ["Acces aux ressources", "Lister les ressources publiques avec pagination", "Citoyen", "Implemente"],
        ["Acces aux ressources", "Filtrer par categorie, type de relation, type de ressource", "Citoyen", "Implemente"],
        ["Acces aux ressources", "Trier par date ou titre", "Citoyen", "Implemente"],
        ["Acces aux ressources", "Afficher le contenu complet d'une ressource", "Citoyen", "Implemente"],
        ["Referentiel", "Lister / filtrer les ressources en back-office", "Administrateur", "Implemente"],
        ["Referentiel", "Ajouter / editer une ressource", "Administrateur", "Implemente"],
        ["Referentiel", "Suspendre une ressource", "Administrateur", "Implemente"],
        ["Referentiel", "Gerer les categories de ressources", "Administrateur", "Implemente"],
        ["Comptes", "Creation d'un compte citoyen", "Citoyen", "Implemente"],
        ["Comptes", "Desactivation / anonymisation de compte (RGPD)", "Citoyen", "Implemente"],
        ["Comptes", "Creation compte moderateur / admin", "Super-admin", "Implemente"],
        ["Creation", "Creation / edition d'une ressource", "Citoyen connecte", "Implemente"],
        ["Creation", "Validation pour publication", "Moderateur", "Implemente"],
        ["Echanges", "Ajout d'un commentaire", "Citoyen connecte", "Implemente"],
        ["Echanges", "Reponse a un commentaire (hierarchie)", "Citoyen connecte", "Implemente"],
        ["Echanges", "Moderation des commentaires", "Moderateur", "Implemente"],
        ["Progression", "Ajouter / retirer des favoris", "Citoyen connecte", "Implemente"],
        ["Progression", "Marquer ressource comme exploitee", "Citoyen connecte", "Implemente"],
        ["Progression", "Mettre de cote une ressource", "Citoyen connecte", "Implemente"],
        ["Progression", "Tableau de bord de progression", "Citoyen connecte", "Implemente"],
        ["Statistiques", "Tableau de bord statistiques", "Administrateur", "Implemente"],
        ["Statistiques", "Filtrer par periode et categorie", "Administrateur", "Implemente"],
        ["Statistiques", "Exporter les statistiques", "Administrateur", "A developper"],
        ["Avance", "Messagerie entre participants", "Citoyen connecte", "A developper"],
        ["Avance", "Invitation de participants", "Citoyen connecte", "A developper"],
        ["Avance", "Traductions multilingues", "Tous", "A developper"],
    ],
    col_widths=[3, 5.5, 3, 2.5]
)

para("Sur les 26 fonctionnalites identifiees, 22 sont pleinement implementees (85%), et 4 restent dans le backlog pour une version ulterieure. Les fonctionnalites implementees couvrent l'integralite du parcours utilisateur principal (consultation, creation, interaction, progression) ainsi que les circuits d'administration et de moderation.")

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 3. CHOIX TECHNOLOGIQUES
# ─────────────────────────────────────────────────────────────────────────────

doc.add_heading("3. Choix technologiques et justifications", level=1)

para("Cette section presente et argumente les choix technologiques retenus pour chaque composant de la plateforme. Chaque decision a ete prise en considerant les contraintes du CCTP, les competences de l'equipe et les bonnes pratiques du marche.")

doc.add_heading("3.1 Backend : Laravel 11 (PHP 8.4)", level=2)
para("Pour le backend, nous avons retenu le framework Laravel dans sa version 11, execute sous PHP 8.4. Ce choix repose sur plusieurs criteres :")

add_table(
    ["Critere", "Laravel", "Alternative (Symfony)", "Alternative (Express.js)"],
    [
        ["Architecture MVC", "Native, convention over config", "Native, plus verbeux", "Non natif, necessite structuration manuelle"],
        ["ORM integre", "Eloquent (Active Record, relations fluides)", "Doctrine (Data Mapper, plus complexe)", "Sequelize/Prisma (tiers)"],
        ["Authentification", "Sanctum integre (tokens API)", "Security Bundle (config complexe)", "Passport.js (tiers)"],
        ["Validation", "FormRequest dedies, regles declaratives", "Validation constraints, annotations", "Joi/Yup (tiers)"],
        ["Ecosysteme", "Riche (migrations, seeders, testing)", "Riche mais courbe plus raide", "Fragmente, choix multiples"],
        ["Courbe d'apprentissage", "Moderee, documentation excellente", "Elevee", "Faible mais manque de structure"],
    ],
    col_widths=[3, 4, 4, 5]
)

para("Laravel s'impose comme le choix optimal pour une API REST avec authentification : son architecture MVC native repond directement a l'exigence du CCTP, Sanctum offre une gestion de tokens legere et securisee, et l'ORM Eloquent simplifie considerablement la modelisation des relations entre entites (ressources, commentaires, favoris, progressions).")

doc.add_heading("3.2 Frontend web : Next.js 16 (React 19, TypeScript)", level=2)
para("Le frontend web repose sur Next.js 16 avec l'App Router, React 19 et TypeScript. Ce choix a ete motive par les elements suivants :")

bullet("App Router : systeme de routage base sur le systeme de fichiers, permettant une organisation claire des pages publiques et administratives via les groupes de routes (main) et administration")
bullet("React 19 : derniere version stable avec des ameliorations de performance (concurrent features) et une gestion native des Server Components")
bullet("TypeScript : typage statique garantissant la fiabilite du code, la detection d'erreurs a la compilation et l'auto-completion IDE")
bullet("Tailwind CSS v4 : framework utilitaire permettant un prototypage rapide et un responsive design natif via les breakpoints (sm, md, lg, xl)")

para("Nous avons ecarte Vue.js/Nuxt.js car l'equipe possedait une meilleure expertise React. Angular a ete ecarte en raison de sa complexite excessive pour un prototype. Le choix de Next.js plutot que Create React App s'explique par les besoins SEO (Server-Side Rendering pour les pages publiques) et la structure de routing plus mature.")

doc.add_heading("3.3 Base de donnees : PostgreSQL 16", level=2)
para("PostgreSQL a ete retenu plutot que MySQL pour les raisons suivantes :")
bullet("Support natif du chiffrement et des extensions de securite, crucial pour notre implementation RGPD")
bullet("Types de donnees riches (text, jsonb) adaptes au stockage de contenus varies")
bullet("Robustesse et conformite ACID assurant l'integrite des donnees relationnelles")
bullet("Compatibilite native avec Laravel via le driver pdo_pgsql")
bullet("Image Docker legere (postgres:16-alpine) pour l'environnement de developpement")

doc.add_heading("3.4 Application mobile : React Native (Expo)", level=2)
para("Pour l'application mobile, le choix s'est porte sur React Native avec le framework Expo, permettant un developpement cross-platform (iOS et Android) a partir d'une base de code unique. Ce choix presente plusieurs avantages :")
bullet("Partage de logique metier avec le frontend web (meme ecosysteme React/TypeScript)")
bullet("Hot reload accelerant le cycle de developpement")
bullet("Expo simplifie la configuration native (camera, notifications, stockage)")
bullet("Large communaute et ecosysteme de composants disponibles")

doc.add_heading("3.5 Infrastructure : Docker Compose", level=2)
para("L'environnement de developpement est entierement containerise via Docker Compose, assurant la reproductibilite entre les machines des developpeurs :")
bullet("PostgreSQL 16 (Alpine) avec health checks et volume persistant")
bullet("API Laravel (PHP 8.4 CLI Alpine) avec auto-migration au demarrage")
bullet("Adminer 4 pour l'administration visuelle de la base de donnees")
bullet("Next.js (Node 22 Alpine) avec hot reload et store pnpm persistant")

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 4. ARCHITECTURE APPLICATIVE
# ─────────────────────────────────────────────────────────────────────────────

doc.add_heading("4. Architecture applicative", level=1)

doc.add_heading("4.1 Architecture globale du systeme", level=2)
para("L'architecture de la plateforme (RE)Sources Relationnelles suit un modele client-serveur decouple. Le backend expose une API REST JSON consommee independamment par le frontend web et l'application mobile. La base de donnees PostgreSQL est accessible uniquement par l'API, garantissant un point d'acces unique et securise aux donnees.")

visual_placeholder(
    "Schema d'architecture globale",
    "4.1 Architecture globale",
    "Schema montrant : Client Web (Next.js) et Client Mobile (React Native/Expo) communiquant via HTTP/JSON avec l'API REST (Laravel). L'API communique avec PostgreSQL. Docker Compose englobe API + DB + Adminer + Web. Fleches : requetes HTTP, reponses JSON, requetes SQL.",
    "draw.io / Lucidchart",
    "Demi-page",
    "Figure 4 — Architecture globale du systeme"
)

para("Cette architecture decouplée offre plusieurs avantages : independance des cycles de deploiement frontend/backend, possibilite de consommer la meme API depuis differents clients (web, mobile, futur backoffice tiers), et testabilite accrue de chaque couche.")

doc.add_heading("4.2 Pattern MVC — Implementation Laravel", level=2)
para("Conformement aux exigences du CCTP, l'architecture backend adopte strictement le pattern Modele-Vue-Controleur (MVC) de Laravel :")

add_table(
    ["Couche", "Responsabilite", "Implementation"],
    [
        ["Modele (Model)", "Representation des donnees, relations, regles metier", "8 modeles Eloquent dans api/app/Models/ : User, Resource, Comment, Category, Favorite, Progression, RelationType, ResourceType"],
        ["Vue (View)", "Presentation des donnees au client", "Reponses JSON structurees (API-first). Les vues HTML sont gerees par le frontend Next.js."],
        ["Controleur (Controller)", "Traitement des requetes, logique applicative", "10 controleurs dans api/app/Http/Controllers/ organises par domaine fonctionnel"],
    ],
    col_widths=[3, 5, 8]
)

para("L'organisation des controleurs par domaine fonctionnel permet une separation claire des responsabilites :")

add_table(
    ["Controleur", "Domaine", "Methodes principales"],
    [
        ["AuthController", "Authentification", "register, login, logout, me, destroy (anonymisation)"],
        ["ResourceController", "Ressources", "index (listing pagine), show (detail), store (creation), update (modification)"],
        ["CommentController", "Commentaires", "indexByResource (arbre), store (creation), reply (reponse)"],
        ["FavoriteController", "Favoris", "store (ajout), destroy (retrait)"],
        ["ProgressionController", "Progression", "exploit (exploiter), setAside (mettre de cote), index (tableau de bord)"],
        ["AdminController", "Administration", "statistics (metriques), suspendResource (suspension)"],
        ["ModerationController", "Moderation", "validateResource, approveComment, deleteComment"],
        ["SuperAdminController", "Super-admin", "createPrivilegedUser (creation comptes privilegies)"],
        ["CategoryController", "Categories", "index (liste des categories)"],
    ],
    col_widths=[3.5, 3, 9]
)

doc.add_heading("4.3 Modele conceptuel de donnees", level=2)

visual_placeholder(
    "Modele Conceptuel de Donnees (MCD Merise)",
    "4.3 Modele conceptuel de donnees",
    "MCD Merise avec les entites : UTILISATEUR, RESSOURCE, COMMENTAIRE, CATEGORIE, TYPE_RELATION, TYPE_RESSOURCE, FAVORI, PROGRESSION. Relations : CREER (1,n), APPARTENIR (1,1), COMMENTER (0,n), AJOUTER_FAVORI (0,n), SUIVRE_PROGRESSION (0,n). Cardinalites precises sur chaque association.",
    "Looping / draw.io / JMerise",
    "Pleine page",
    "Figure 5 — Modele Conceptuel de Donnees (MCD Merise)"
)

para("Le modele de donnees repose sur 8 tables principales, concu pour couvrir l'ensemble des fonctionnalites metier. Les contraintes d'integrite garantissent la coherence des donnees :")

add_table(
    ["Table", "Role metier", "Colonnes principales"],
    [
        ["users", "Comptes utilisateurs avec donnees chiffrees", "id, name (chiffre), email (chiffre), email_hash (SHA-256), password (hashe), role (enum), is_active"],
        ["resources", "Ressources publiees ou en attente de validation", "id, title, content, user_id (FK), category_id (FK), relation_type_id (FK), resource_type_id (FK), status (enum), is_public"],
        ["comments", "Commentaires hierarchiques avec moderation", "id, content, user_id (FK), resource_id (FK), parent_id (FK self), is_approved"],
        ["favorites", "Relation utilisateur-ressource pour les favoris", "id, user_id (FK), resource_id (FK) — unique(user_id, resource_id)"],
        ["progressions", "Etat d'avancement utilisateur-ressource", "id, user_id (FK), resource_id (FK), status (enum: exploited/set_aside) — unique(user_id, resource_id)"],
        ["categories", "Categories de ressources", "id, name, description"],
        ["relation_types", "Types de relations ciblees (couple, famille, collegues...)", "id, name"],
        ["resource_types", "Types de contenus (article, video, activite...)", "id, name"],
    ],
    col_widths=[2.5, 4, 9.5]
)

doc.add_heading("4.4 Diagramme de classes UML", level=2)

visual_placeholder(
    "Diagramme de classes UML",
    "4.4 Diagramme de classes",
    "Diagramme de classes UML montrant les 8 entites avec leurs attributs types, leurs methodes publiques et les associations : User 1-->* Resource, User 1-->* Comment, User 1-->* Favorite, User 1-->* Progression, Resource *-->1 Category, Resource *-->1 RelationType, Resource *-->1 ResourceType, Resource 1-->* Comment, Comment 0..1-->* Comment (auto-reference replies). Inclure les multiplicites et la navigabilite.",
    "StarUML / draw.io / PlantUML",
    "Pleine page",
    "Figure 6 — Diagramme de classes UML"
)

para("Les relations Eloquent implementees dans les modeles Laravel refletent directement ce diagramme :")
bullet("User -> hasMany(Resource, Comment, Favorite, Progression)")
bullet("Resource -> belongsTo(User, Category, RelationType, ResourceType) et hasMany(Comment, Favorite, Progression)")
bullet("Comment -> belongsTo(User, Resource, Comment:parent) et hasMany(Comment:replies)")

para("Les contraintes d'integrite garantissent la coherence des donnees :")
bullet("Cles etrangeres avec CASCADE sur toutes les relations metier (suppression en cascade)")
bullet("Contrainte d'unicite sur favorites(user_id, resource_id) pour empecher les doublons de favoris")
bullet("Contrainte d'unicite sur progressions(user_id, resource_id) pour un etat unique par ressource/utilisateur")
bullet("Index sur resources.status et resources.is_public pour optimiser les requetes de listing")
bullet("Index sur comments.is_approved pour filtrer rapidement les commentaires moderes")
bullet("Colonne users.email_hash (SHA-256) unique pour garantir l'unicite fonctionnelle malgre le chiffrement de l'email")
bullet("Favorite et Progression servent de tables pivot enrichies avec contrainte d'unicite (user_id, resource_id)")

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 5. DEVELOPPEMENT BACKEND
# ─────────────────────────────────────────────────────────────────────────────

doc.add_heading("5. Developpement backend — API REST Laravel", level=1)

doc.add_heading("5.1 Organisation des routes API", level=2)
para("L'API REST est structuree autour de groupes de routes proteges par des middlewares progressifs. Toutes les reponses sont forcees en JSON.")

add_table(
    ["Groupe", "Middleware", "Routes principales"],
    [
        ["Public", "Aucun (+ throttle:auth sur register/login)", "GET /resources, GET /resources/{id}, GET /categories, GET /resources/{id}/comments, POST /register, POST /login"],
        ["Authentifie", "auth:sanctum", "POST /logout, GET /user, DELETE /user, POST /resources, PUT /resources/{id}, POST /resources/{id}/comments, POST /comments/{id}/reply, favoris, progression"],
        ["Administration", "auth:sanctum + role:admin,super_admin", "GET /admin/statistics, PUT /admin/resources/{id}/suspend"],
        ["Moderation", "auth:sanctum + role:moderator,admin,super_admin", "PUT /moderation/resources/{id}/validate, PUT /moderation/comments/{id}/approve, DELETE /moderation/comments/{id}"],
        ["Super-admin", "auth:sanctum + role:super_admin", "POST /super-admin/users"],
    ],
    col_widths=[2.5, 5, 8.5]
)

para("Le tableau suivant detaille l'ensemble des endpoints implementes avec leur methode HTTP, leur URI, leur controleur et la reponse attendue :")

add_table(
    ["Methode", "Endpoint", "Controleur / Action", "Reponse"],
    [
        ["POST", "/api/register", "AuthController@register", "201 — token + user"],
        ["POST", "/api/login", "AuthController@login", "200 — token + user"],
        ["POST", "/api/logout", "AuthController@logout", "200 — message"],
        ["GET", "/api/user", "AuthController@me", "200 — user"],
        ["DELETE", "/api/user", "AuthController@destroy", "200 — anonymise"],
        ["GET", "/api/resources", "ResourceController@index", "200 — pagine (15/page)"],
        ["GET", "/api/resources/{id}", "ResourceController@show", "200 — resource + relations"],
        ["POST", "/api/resources", "ResourceController@store", "201 — resource (PENDING)"],
        ["PUT", "/api/resources/{id}", "ResourceController@update", "200 — resource MAJ"],
        ["GET", "/api/categories", "CategoryController@index", "200 — categories[]"],
        ["GET", "/api/resources/{id}/comments", "CommentController@indexByResource", "200 — arbre commentaires"],
        ["POST", "/api/resources/{id}/comments", "CommentController@store", "201 — commentaire"],
        ["POST", "/api/comments/{id}/reply", "CommentController@reply", "201 — reponse"],
        ["POST", "/api/resources/{id}/favorite", "FavoriteController@store", "201 ou 200"],
        ["DELETE", "/api/resources/{id}/favorite", "FavoriteController@destroy", "200"],
        ["POST", "/api/resources/{id}/exploit", "ProgressionController@exploit", "200/201"],
        ["POST", "/api/resources/{id}/set-aside", "ProgressionController@setAside", "200/201"],
        ["GET", "/api/progression", "ProgressionController@index", "200 — dashboard"],
        ["GET", "/api/admin/statistics", "AdminController@statistics", "200 — metriques"],
        ["PUT", "/api/admin/resources/{id}/suspend", "AdminController@suspendResource", "200"],
        ["PUT", "/api/moderation/resources/{id}/validate", "ModerationController@validateResource", "200"],
        ["PUT", "/api/moderation/comments/{id}/approve", "ModerationController@approveComment", "200"],
        ["DELETE", "/api/moderation/comments/{id}", "ModerationController@deleteComment", "200"],
        ["POST", "/api/super-admin/users", "SuperAdminController@createPrivilegedUser", "201 — user"],
    ],
    col_widths=[1.5, 5, 5, 3.5]
)

doc.add_heading("5.2 Validation et sanitization des entrees", level=2)
para("Chaque endpoint d'ecriture utilise un FormRequest dedie qui assure la validation et la sanitization des donnees avant traitement. Un trait SanitizesInput centralise les operations de nettoyage :")

bullet("strip_tags() : suppression de toute balise HTML pour prevenir les attaques XSS")
bullet("Normalisation des espaces : remplacement des espaces multiples par un espace unique")
bullet("trim() : suppression des espaces en debut et fin de chaine")
bullet("strtolower() + trim() sur les emails : normalisation pour garantir l'unicite")
bullet("Hashage SHA-256 de l'email : generation d'un email_hash pour les recherches sans dechiffrement")

para("Exemples de regles de validation implementees :")

add_table(
    ["FormRequest", "Champ", "Regles"],
    [
        ["RegisterRequest", "name", "required, string, min:2, max:255"],
        ["RegisterRequest", "email", "required, string, email:rfc, max:255"],
        ["RegisterRequest", "email_hash", "required, string, size:64, unique:users"],
        ["RegisterRequest", "password", "required, confirmed, Password::defaults()"],
        ["StoreResourceRequest", "title", "required, string, min:3, max:255"],
        ["StoreResourceRequest", "content", "required, string, min:10"],
        ["StoreResourceRequest", "category_id", "required, integer, exists:categories,id"],
        ["CommentRequest", "content", "required, string, min:2, max:2000"],
    ],
    col_widths=[3.5, 3, 9.5]
)

doc.add_heading("5.3 Flux d'authentification", level=2)
para("L'authentification repose sur Laravel Sanctum avec des tokens Bearer. Le flux complet est le suivant :")

visual_placeholder(
    "Diagramme de sequence — Authentification",
    "5.3 Flux d'authentification",
    "Diagramme de sequence UML montrant le flux : Client -> API (POST /login avec email + password) -> API recherche via email_hash dans PostgreSQL -> Verification Hash::check(password) -> Verification is_active -> Creation token Sanctum -> Retour token Bearer + user au client. Puis : Client -> API (GET /user + Authorization: Bearer token) -> Validation Sanctum -> Retour user authentifie.",
    "PlantUML / draw.io / Mermaid",
    "Demi-page",
    "Figure 7 — Diagramme de sequence du flux d'authentification"
)

para("Points cles de l'implementation :")
bullet("L'email est stocke chiffre en base ; la recherche se fait via email_hash (SHA-256 de l'email normalise)")
bullet("Le mot de passe est verifie via Hash::check() apres recuperation de l'utilisateur")
bullet("Le flag is_active est verifie pour bloquer les comptes desactives ou anonymises")
bullet("Un token Sanctum est cree via createToken('auth_token') et retourne au client")
bullet("Le token est transmis dans l'en-tete Authorization: Bearer <token> pour chaque requete authentifiee")
bullet("La deconnexion revoque le token courant via currentAccessToken()->delete()")

doc.add_heading("5.4 Controle d'acces par roles", level=2)
para("Le controle d'acces repose sur deux mecanismes complementaires :")

para("RoleMiddleware : middleware dynamique acceptant une liste de roles autorises (ex : role:admin,super_admin). Il verifie que le role de l'utilisateur authentifie figure dans la liste autorisee et retourne une erreur 403 en cas de refus.", bold=False)

para("ResourcePolicy : policy Laravel garantissant que seul l'auteur d'une ressource peut la modifier. L'autorisation est verifiee via $this->authorize('update', $resource) dans le controleur.", bold=False)

add_table(
    ["Role", "Enum", "Droits"],
    [
        ["Citoyen", "citizen", "Consultation, creation de ressources (PENDING), commentaires, favoris, progression"],
        ["Moderateur", "moderator", "Validation de ressources (PENDING -> PUBLISHED), approbation/suppression de commentaires"],
        ["Administrateur", "admin", "Statistiques, suspension de ressources (-> ARCHIVED), droits moderateur inclus"],
        ["Super-admin", "super_admin", "Creation de comptes privilegies (moderateur, admin), tous les droits"],
    ],
    col_widths=[3, 2.5, 10.5]
)

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 6. SECURITE ET RGPD
# ─────────────────────────────────────────────────────────────────────────────

doc.add_heading("6. Securite et conformite RGPD", level=1)

doc.add_heading("6.1 Mesures de securite implementees", level=2)
para("La securite du backend repose sur plusieurs couches complementaires, conformement aux exigences du CCTP en matiere de chiffrement, d'anonymisation et de protection des donnees sensibles.")

add_table(
    ["Mesure", "Implementation technique", "Fichier source"],
    [
        ["Hashage mot de passe", "Cast Laravel 'hashed' sur User.password (bcrypt)", "api/app/Models/User.php"],
        ["Chiffrement donnees personnelles", "Crypt::encryptString() sur name et email", "api/app/Models/User.php (mutators)"],
        ["Lookup email securise", "email_hash = SHA-256(email normalise) unique en base", "api/app/Models/User.php + migration"],
        ["Rate limiting", "throttle:auth sur /register et /login (5 tentatives/min)", "api/routes/api.php"],
        ["Headers de securite", "Middleware SecurityHeaders (CSP, X-Frame-Options, etc.)", "api/app/Http/Middleware/SecurityHeaders.php"],
        ["Sanitization des entrees", "Trait SanitizesInput (strip_tags, trim, normalisation)", "api/app/Http/Requests/Concerns/SanitizesInput.php"],
        ["Controle d'acces", "RoleMiddleware + ResourcePolicy", "api/app/Http/Middleware/RoleMiddleware.php"],
        ["Protection CSRF", "Token Sanctum Bearer (pas de cookies de session)", "Configuration Sanctum"],
    ],
    col_widths=[3.5, 6, 6.5]
)

doc.add_heading("6.2 Headers de securite HTTP", level=2)
para("Le middleware SecurityHeaders ajoute les en-tetes suivants a chaque reponse API :")
bullet("X-Content-Type-Options: nosniff — empeche le navigateur d'interpreter le type MIME")
bullet("X-Frame-Options: DENY — bloque l'inclusion de l'API dans une iframe (protection clickjacking)")
bullet("Referrer-Policy: no-referrer — empeche la fuite d'informations via le header Referer")
bullet("X-XSS-Protection: 1; mode=block — active le filtre XSS integre du navigateur")
bullet("Content-Security-Policy: default-src 'none'; frame-ancestors 'none'; base-uri 'self'; form-action 'self' — politique de securite du contenu restrictive")

doc.add_heading("6.3 Conformite RGPD — Chiffrement et anonymisation", level=2)
para("La conformite RGPD a ete implementee selon deux axes majeurs :")

para("Chiffrement des donnees personnelles a la source :", bold=True)
para("Les champs name et email du modele User sont automatiquement chiffres a l'ecriture via Crypt::encryptString() et dechiffres a la lecture via Crypt::decryptString(). Ce mecanisme est transparent grace aux accesseurs/mutateurs Eloquent. Un champ email_hash (SHA-256 de l'email normalise en minuscules) permet les recherches d'authentification sans dechiffrer l'ensemble des emails en base.")

para("Droit a l'effacement — Anonymisation logique :", bold=True)
para("Plutot qu'une suppression physique du compte (qui briserait l'integrite referentielle des ressources, commentaires et favoris), l'endpoint DELETE /api/user applique une anonymisation logique via la methode User::anonymize() :")
bullet("Revocation de tous les tokens d'acces actifs")
bullet("Remplacement du nom par 'Deleted User'")
bullet("Remplacement de l'email par une adresse technique deleted-user-{id}-{uuid}@example.invalid")
bullet("Remise du role en citizen et desactivation du compte (is_active = false)")
bullet("Effacement du mot de passe et de la date de verification email")

para("Ce choix preserve la coherence historique des contenus tout en garantissant l'impossibilite d'identifier l'ancien utilisateur.")

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 7. DEVELOPPEMENT FRONTEND WEB
# ─────────────────────────────────────────────────────────────────────────────

doc.add_heading("7. Developpement frontend web — Next.js", level=1)

doc.add_heading("7.1 Architecture de l'application web", level=2)
para("L'application web est construite avec Next.js 16 (App Router) et organisee en deux espaces distincts :")
bullet("Espace public (main) : pages accessibles a tous les visiteurs (accueil, ressources, aide, presentation, authentification)")
bullet("Espace administration : pages reservees aux roles privilegies (tableau de bord, statistiques, moderation, gestion des utilisateurs et ressources)")

add_table(
    ["Route", "Page", "Description"],
    [
        ["/", "Accueil", "Hero, categories, fonctionnement, CTA — vitrine de la plateforme"],
        ["/ressources", "Listing ressources", "Recherche, filtres (categorie, type relation, type ressource), tri, pagination"],
        ["/ressources/[id]", "Detail ressource", "Contenu complet, commentaires hierarchiques, ajout favori, formulaire commentaire"],
        ["/auth/connexion", "Connexion", "Formulaire email/mot de passe avec gestion d'erreurs"],
        ["/auth/inscription", "Inscription", "Formulaire complet avec validation champ par champ et confirmation mot de passe"],
        ["/presentation", "Presentation", "Page d'information sur la mission de la plateforme"],
        ["/aide", "Aide / FAQ", "Centre d'aide avec questions frequentes"],
        ["/administration", "Accueil admin", "Page d'accueil de l'espace d'administration"],
        ["/administration/stats", "Statistiques", "Tableau de bord avec graphiques (barres, donut), metriques cles, table utilisateurs"],
        ["/administration/moderation", "Moderation", "Onglets ressources en attente / commentaires a moderer"],
        ["/administration/ressources", "Gestion ressources", "CRUD ressources avec edition/suppression"],
        ["/administration/utilisateur", "Gestion utilisateurs", "Table des utilisateurs, creation, edition, suppression"],
        ["/administration/categorie", "Gestion categories", "CRUD categories de ressources"],
    ],
    col_widths=[3.5, 3, 9.5]
)

doc.add_heading("7.2 Composants UI reutilisables", level=2)
para("Une bibliotheque de composants UI a ete developpee pour assurer la coherence visuelle et l'accessibilite sur l'ensemble de l'application :")

add_table(
    ["Composant", "Variantes", "Accessibilite"],
    [
        ["Button", "primary, secondary, accent, outline — tailles sm/md/lg", "focus-visible:ring-2, disabled:opacity-50"],
        ["Input", "Avec label et erreur integres", "aria-invalid, aria-describedby, role='alert' sur les erreurs"],
        ["Card", "Padding none/sm/md/lg, hover optionnel", "Structure semantique, transition fluide"],
        ["Badge", "primary, secondary, accent, gray, success, warning, error", "Texte descriptif (pas de dependance a la couleur seule)"],
    ],
    col_widths=[2.5, 6, 7.5]
)

doc.add_heading("7.3 Gestion de l'authentification cote client", level=2)
para("L'authentification est geree via un React Context (AuthContext) qui encapsule l'etat utilisateur et les methodes d'authentification. Le token Sanctum est stocke dans le localStorage du navigateur et transmis automatiquement dans les en-tetes de chaque requete API via un client HTTP centralise (api.ts).")

para("Le flux cote client est le suivant :")
bullet("Au chargement de l'application, le contexte verifie la presence d'un token en localStorage")
bullet("Si un token est present, une requete GET /api/user est envoyee pour rehydrater l'etat utilisateur")
bullet("En cas d'echec (token expire ou invalide), le token est supprime du localStorage")
bullet("A la connexion/inscription, le token retourne par l'API est stocke et l'utilisateur redirige")
bullet("A la deconnexion, le token est revoque cote serveur (POST /api/logout) puis supprime du localStorage")

doc.add_heading("7.4 Charte graphique et palette de couleurs", level=2)
para("La palette de couleurs est definie via des variables CSS Tailwind personnalisees :")

add_table(
    ["Variable", "Couleur", "Usage"],
    [
        ["--primary", "#1D3A6D (bleu fonce)", "Couleur principale, navigation, titres, boutons principaux"],
        ["--primary-light", "#2D4A7D", "Variante hover sur les elements primaires"],
        ["--primary-dark", "#0D2A5D", "Variante foncee pour les contrastes"],
        ["--secondary", "#4CAF50 (vert)", "Actions secondaires, validations, succes"],
        ["--accent", "#FFC107 (jaune/ambre)", "Mise en avant, call-to-action, elements decoratifs"],
        ["Gris (50-900)", "Echelle de gris complete", "Textes, bordures, fonds, separateurs"],
    ],
    col_widths=[3, 4, 9]
)

para("L'identite visuelle s'appuie sur la police Inter (Google Fonts) pour le corps de texte, avec une approche responsive-first. Les composants de l'espace administration utilisent des design tokens dedies (Poppins pour les titres, Open Sans pour le corps) afin de differencier visuellement l'espace public de l'espace d'administration.")

doc.add_heading("7.5 Espace d'administration et tableau de bord", level=2)
para("L'espace d'administration constitue une partie majeure de l'application web, accessible aux roles admin et super_admin. Il offre une vue complete sur l'activite de la plateforme et des outils de gestion.")

visual_placeholder(
    "Captures d'ecran de l'interface d'administration",
    "7.5 Espace d'administration",
    "Captures montrant : (1) Le tableau de bord statistiques avec les graphiques en barres et en donut, les cartes de metriques (total utilisateurs, utilisateurs actifs, nouveaux ce mois, desactives). (2) La page de moderation avec les onglets ressources/commentaires. (3) La page de gestion des utilisateurs avec le tableau et les actions.",
    "Captures d'ecran du navigateur",
    "Pleine page (3 captures)",
    "Figure 9 — Interface d'administration : tableau de bord, moderation et gestion des utilisateurs"
)

para("Le tableau de bord statistiques (route /administration/stats) affiche :")
bullet("Metriques cles : nombre total d'utilisateurs, utilisateurs actifs, nouveaux inscrits ce mois, comptes desactives")
bullet("Graphique en barres : evolution des inscriptions par mois")
bullet("Graphique en donut : repartition des ressources par categorie")
bullet("Table des utilisateurs deconnectes/desactives")
bullet("Panel d'administration avec actions rapides")

para("La page de moderation (route /administration/moderation) propose deux onglets :")
bullet("Onglet Ressources : cartes de ressources en attente de validation avec boutons Approuver/Refuser et contenu deroulant")
bullet("Onglet Commentaires : liste des commentaires avec filtrage par signalement, actions d'edition et de suppression, modales de confirmation")

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 8. ACCESSIBILITE RGAA
# ─────────────────────────────────────────────────────────────────────────────

doc.add_heading("8. Accessibilite — Conformite RGAA", level=1)

para("Le Referentiel General d'Amelioration de l'Accessibilite (RGAA) constitue une exigence explicite du CCTP. L'application web a ete developpee en integrant les bonnes pratiques d'accessibilite des le debut du developpement (approche 'accessibility-first').")

doc.add_heading("8.1 Mesures implementees", level=2)

add_table(
    ["Critere RGAA", "Implementation", "Fichier"],
    [
        ["Navigation au clavier", "Lien d'evitement 'Aller au contenu principal' (.skip-link) visible au focus", "apps/web/src/app/(main)/layout.tsx"],
        ["Landmarks ARIA", "nav[aria-label], main[id=main-content], footer — structure semantique complete", "Navbar.tsx, layout.tsx, Footer.tsx"],
        ["Langue de la page", "html[lang='fr'] declare sur le document racine", "apps/web/src/app/layout.tsx"],
        ["Labels de formulaire", "Chaque input associe a un label via htmlFor + id auto-genere", "apps/web/src/components/ui/Input.tsx"],
        ["Etats d'erreur", "aria-invalid, aria-describedby, role='alert' aria-live='polite'", "Input.tsx, inscription/page.tsx"],
        ["Menu responsive", "aria-expanded, aria-controls liant le bouton au menu mobile", "Navbar.tsx"],
        ["Elements decoratifs", "aria-hidden='true' sur les emojis decoratifs", "page.tsx (accueil)"],
        ["Contraste des couleurs", "Primary #1D3A6D sur blanc : ratio > 7:1 (AAA)", "globals.css"],
        ["Focus visible", "outline: 3px solid var(--primary) avec offset 2px sur a, button, input", "globals.css"],
        ["Liens de reseaux sociaux", "aria-label='Facebook/Twitter/LinkedIn' + title descriptif", "Footer.tsx"],
        ["Pagination", "aria-label='Aller a la page N', aria-current='page' sur la page active", "ressources/page.tsx"],
        ["Filtres", "aria-label descriptif sur chaque select de filtrage", "ressources/page.tsx"],
    ],
    col_widths=[3, 6.5, 6.5]
)

doc.add_heading("8.2 Points d'amelioration identifies", level=2)
para("Malgre les efforts d'accessibilite, certains points restent a ameliorer pour atteindre une conformite complete :")
bullet("Ajout systematique d'attributs alt sur les images et icones non decoratives")
bullet("Verification de l'attribut scope='col' sur les en-tetes de tableaux dans les pages d'administration")
bullet("Enrichissement des aria-describedby sur les champs de formulaire complexes")
bullet("Tests automatises avec axe-core ou Lighthouse pour mesurer le score d'accessibilite")
bullet("Declaration formelle du niveau de conformite RGAA vise (AA)")

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 9. APPLICATION MOBILE
# ─────────────────────────────────────────────────────────────────────────────

doc.add_heading("9. Application mobile — React Native / Expo", level=1)

doc.add_heading("9.1 Architecture prevue", level=2)
para("L'application mobile est concue avec React Native et le framework Expo, permettant un deploiement cross-platform iOS et Android. L'architecture reprend les memes principes que l'application web : consommation de l'API REST Laravel via des services HTTP partages.")

visual_placeholder(
    "Schema d'architecture de l'application mobile",
    "9.1 Architecture mobile",
    "Schema montrant : ecrans principaux (Accueil, Liste ressources, Detail ressource, Profil, Connexion) avec navigation (React Navigation stack/tab). Communication avec l'API Laravel via le meme client HTTP. Stockage local du token via AsyncStorage.",
    "draw.io / Figma",
    "Demi-page",
    "Figure 8 — Architecture de l'application mobile"
)

para("L'application mobile est structuree autour des ecrans suivants, conformement aux fonctionnalites du CCTP :")

add_table(
    ["Ecran", "Fonctionnalite", "Correspondance API"],
    [
        ["Accueil", "Presentation et acces rapide aux ressources", "GET /api/resources"],
        ["Liste des ressources", "Listing avec filtres et recherche", "GET /api/resources?category=&sort="],
        ["Detail ressource", "Contenu, commentaires, ajout favori", "GET /api/resources/{id}, GET /api/resources/{id}/comments"],
        ["Connexion / Inscription", "Formulaires d'authentification", "POST /api/login, POST /api/register"],
        ["Profil", "Informations utilisateur, deconnexion", "GET /api/user, POST /api/logout"],
        ["Progression", "Tableau de bord des favoris et exploitations", "GET /api/progression"],
    ],
    col_widths=[3.5, 5, 7.5]
)

doc.add_heading("9.2 Specificites de la plateforme mobile", level=2)
para("Conformement aux exigences du CCTP concernant le 'Mobile First' et les contraintes specifiques au mobile, l'application a ete concue avec les principes suivants :")
bullet("Navigation par onglets (Tab Navigation) pour un acces rapide aux fonctionnalites principales")
bullet("Composants natifs (TouchableOpacity, FlatList, ScrollView) pour des performances optimales")
bullet("AsyncStorage pour le stockage persistant du token d'authentification")
bullet("Adaptation de l'interface aux differentes tailles d'ecran via Dimensions et Flexbox natif")
bullet("Gestion du mode hors-ligne partiel (cache des ressources consultees)")

doc.add_heading("9.3 Responsive design et adaptation", level=2)
para("L'application web est egalement responsive grace a l'utilisation des breakpoints Tailwind CSS : sm (640px), md (768px), lg (1024px), xl (1280px). La navigation passe d'un menu horizontal a un menu burger sur les ecrans mobiles (< 768px) avec une gestion aria-expanded pour l'accessibilite.")

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 10. PLAN DE TESTS
# ─────────────────────────────────────────────────────────────────────────────

doc.add_heading("10. Plan de tests", level=1)

doc.add_heading("10.1 Methodologie de test", level=2)
para("La strategie de test adoptee pour le projet s'articule autour de trois niveaux complementaires :")

add_table(
    ["Niveau", "Type", "Outil", "Couverture"],
    [
        ["Backend — Tests feature", "Tests d'integration API", "PHPUnit + Laravel TestCase", "Endpoints auth, CRUD ressources, interactions, admin, moderation"],
        ["Backend — Tests unitaires", "Tests des modeles et services isoles", "PHPUnit", "Modeles, enums, sanitization"],
        ["Frontend — Tests manuels", "Verification des parcours utilisateur", "Navigateur + outils developpeur", "Pages publiques, admin, moderation, formulaires"],
        ["Accessibilite", "Audit RGAA semi-automatise", "Lighthouse / axe DevTools", "Contraste, landmarks, aria, navigation clavier"],
    ],
    col_widths=[3.5, 4, 4, 4.5]
)

para("La base de donnees de test utilise le trait RefreshDatabase de Laravel, qui reinitialise la base entre chaque test pour garantir l'isolation. Les tokens Sanctum sont generes dans les tests via actingAs() ou createToken().")

doc.add_heading("10.2 Cahier de tests backend", level=2)
para("Les tests feature backend couvrent l'ensemble des cas metier critiques. Voici le detail des suites de tests implementees :")

para("Suite AuthEndpointsTest (10 tests) :", bold=True)
add_table(
    ["Test", "Scenario", "Verification"],
    [
        ["register_creates_citizen", "Inscription d'un nouvel utilisateur", "Token retourne, role citizen, donnees sanitizees et chiffrees en base"],
        ["login_returns_token", "Connexion avec identifiants valides", "Token Sanctum retourne, recherche via email_hash"],
        ["login_rate_limited", "5+ tentatives de login echouees", "HTTP 429 (Too Many Requests) apres 5 echecs"],
        ["register_rate_limited", "5+ tentatives d'inscription", "HTTP 429 apres 5 tentatives"],
        ["logout_revokes_token", "Deconnexion utilisateur", "Token supprime de personal_access_tokens"],
        ["user_returns_auth", "Acces au profil authentifie", "Donnees utilisateur retournees correctement"],
        ["protected_routes_require_auth", "Acces sans token", "HTTP 401 sur toutes les routes protegees"],
        ["deletion_anonymizes", "Suppression de compte", "Nom='Deleted User', email anonymise, tokens revoques"],
        ["role_middleware", "Acces avec mauvais role", "HTTP 403 pour non-admin sur routes admin"],
        ["security_headers", "Headers de securite", "Presence de CSP, X-Frame-Options, etc."],
    ],
    col_widths=[3, 5, 8]
)

para("Suite ResourceEndpointsTest (6 tests) :", bold=True)
add_table(
    ["Test", "Scenario", "Verification"],
    [
        ["index_paginated_filtered", "Listing pagine avec filtres", "15 resultats/page, filtres categorie/type, tri par date/titre"],
        ["index_sort_pagination", "Tri et pagination combinees", "Ordre correct, navigation entre pages"],
        ["show_resource_detail", "Affichage detail ressource", "Donnees completes avec relations chargees"],
        ["store_pending_resource", "Creation par citoyen", "Statut PENDING, donnees sanitizees"],
        ["update_author_only", "Modification par auteur vs autre", "Auteur autorise, autre utilisateur refuse (403)"],
        ["categories_list", "Liste des categories", "Retour ordonne par nom"],
    ],
    col_widths=[3.5, 5, 7.5]
)

para("Suite InteractionEndpointsTest (6 tests) :", bold=True)
bullet("Listing des commentaires approuves avec arbre de reponses")
bullet("Creation de commentaire par un citoyen (avec sanitization)")
bullet("Refus de creation pour un non-citoyen (role enforcement)")
bullet("Reponse a un commentaire existant (parent_id)")
bullet("Ajout et retrait de favoris (cycle complet)")
bullet("Progression : exploitation, mise de cote et tableau de bord")

para("Suite AdminEndpointsTest (5 tests) :", bold=True)
bullet("Statistiques admin avec filtres par periode et categorie")
bullet("Suspension de ressource (statut ARCHIVED)")
bullet("Validation de ressource et moderation de commentaires")
bullet("Creation de compte privilegie par super-admin")
bullet("Protection des routes admin par role")

doc.add_heading("10.3 Recette et PV", level=2)
para("Le cahier de recettes et les proces-verbaux de tests sont formalises dans les documents annexes :")
bullet("Cahier_de_Recettes.pdf : scenarios de recette fonctionnelle couvrant les parcours utilisateur complets")
bullet("Cahier_de_tests.pdf : matrice de tests detaillee avec cas nominaux, cas limites et cas d'erreur")

para("Ces documents sont fournis en annexe du present dossier technique.")

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 11. DEPLOIEMENT
# ─────────────────────────────────────────────────────────────────────────────

doc.add_heading("11. Deploiement et environnement technique", level=1)

doc.add_heading("11.1 Architecture Docker Compose", level=2)
para("L'ensemble de l'environnement applicatif est containerise via Docker Compose, garantissant la reproductibilite entre les postes de developpement et la coherence de l'environnement de test.")

add_table(
    ["Service", "Image", "Port", "Role"],
    [
        ["db", "postgres:16-alpine", "5432", "Base de donnees PostgreSQL avec healthcheck et volume persistant"],
        ["api", "php:8.4-cli-alpine (Dockerfile custom)", "8000", "API Laravel avec migrations automatiques au demarrage"],
        ["db_admin", "adminer:4", "8080", "Interface web d'administration de la BDD"],
        ["web", "node:22-alpine", "3005 (configurable)", "Frontend Next.js avec hot reload et pnpm"],
    ],
    col_widths=[2, 5, 1.5, 7.5]
)

doc.add_heading("11.2 Initialisation automatique", level=2)
para("Au demarrage du conteneur API, les etapes suivantes sont executees automatiquement :")
bullet("Copie du fichier .env.example en .env si absent")
bullet("Installation des dependances Composer (composer install)")
bullet("Generation de la cle d'application Laravel (php artisan key:generate)")
bullet("Execution des migrations de base de donnees (php artisan migrate --force)")
bullet("Lancement du serveur de developpement (php artisan serve --host=0.0.0.0)")

doc.add_heading("11.3 Donnees de demonstration", level=2)
para("Le seeding Laravel genere un jeu de donnees coherent pour les tests et demonstrations :")
bullet("10 utilisateurs de demonstration couvrant tous les roles (super_admin, admin, moderator, citizen)")
bullet("20 ressources avec differents statuts (published, pending)")
bullet("15 commentaires avec des reponses hierarchiques")
para("Mot de passe de demonstration pour tous les comptes : password123")

doc.add_heading("11.4 Variables d'environnement", level=2)
add_table(
    ["Variable", "Valeur par defaut", "Description"],
    [
        ["DB_CONNECTION", "pgsql", "Driver de base de donnees"],
        ["DB_HOST", "db", "Hote PostgreSQL (nom du service Docker)"],
        ["DB_DATABASE", "ressources_relationnelles", "Nom de la base de donnees"],
        ["CORS_ALLOWED_ORIGINS", "localhost:3000,3005,19006,8081", "Origines autorisees (web, Expo)"],
        ["NEXT_PUBLIC_API_URL", "http://localhost:8000", "URL de l'API pour le frontend"],
        ["WEB_PORT", "3005", "Port du frontend Next.js"],
    ],
    col_widths=[4, 4, 8]
)

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 12. ORGANISATION DU PROJET
# ─────────────────────────────────────────────────────────────────────────────

doc.add_heading("12. Organisation du projet", level=1)

doc.add_heading("12.1 Equipe et repartition des roles", level=2)

add_table(
    ["Membre", "GitHub", "Role", "Perimetre technique"],
    [
        ["Chamil MUSAYEV", "@musayevchamil", "Backend Lead", "API Laravel, BDD PostgreSQL, Auth Sanctum, Securite, RGPD"],
        ["Jamal", "@jamal09102007", "Frontend Lead", "Next.js, charte graphique, accessibilite RGAA, integration API"],
        ["Romain CONDOMITTI", "@romain-condomitti", "Back-office", "Admin, Auth frontend, integration composants, moderation"],
        ["Leon HEU", "@leonheu", "Mobile Lead", "React Native, Expo, screens multi-plateformes"],
    ],
    col_widths=[3, 3.5, 2.5, 7]
)

doc.add_heading("12.2 Workflow Git et methodologie", level=2)
para("L'equipe a adopte un workflow Git Flow simplifie avec deux branches principales :")

add_table(
    ["Branche", "Role", "Politique de merge"],
    [
        ["main", "Production — code stable, pret a livrer", "Merge de develop en fin de sprint uniquement"],
        ["develop", "Integration — toutes les features mergees", "Via Pull Request avec review obligatoire"],
        ["feature/*", "Branches de developpement par fonctionnalite", "Squash and merge vers develop apres approbation"],
    ],
    col_widths=[3, 5, 8]
)

para("Regles imposees : aucun push direct sur main ou develop ; tout passe par une Pull Request avec au moins un review par un autre developpeur ; merge en Squash and merge pour un historique propre.")

doc.add_heading("12.3 Conventions de commits", level=2)
para("L'equipe utilise les Conventional Commits pour structurer l'historique :")
bullet("feat: nouvelle fonctionnalite (ex: feat: ajouter endpoint CRUD ressources)")
bullet("fix: correction de bug (ex: fix: corriger le token expire au login)")
bullet("docs: documentation (ex: docs: ajouter le chapitre architecture)")
bullet("test: ajout ou modification de tests (ex: test: ajouter tests unitaires auth)")
bullet("refactor: refactorisation sans changement fonctionnel")
bullet("style: modifications visuelles (ex: style: ajuster les couleurs du header)")
bullet("chore: maintenance (ex: chore: configurer eslint + prettier)")

doc.add_heading("12.4 Sprints et jalons", level=2)

add_table(
    ["Sprint", "Nom", "Periode", "Objectif"],
    [
        ["S1", "Fondations", "Semaine du 10 mars", "Setup projets (Laravel, Next.js, Expo), authentification, navigation de base"],
        ["S2", "Fonctionnalites coeur", "Semaine du 17 mars", "CRUD ressources, commentaires, favoris, integration API"],
        ["S3", "Fonctionnalites avancees", "Semaine du 24 mars", "Admin, statistiques, RGAA, progression utilisateur"],
        ["S4", "Tests, polish & soutenance", "Semaine du 31 mars", "Tests, documentation technique, corrections, preparation soutenance"],
    ],
    col_widths=[1.5, 3.5, 3.5, 7.5]
)

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 13. CONCLUSION
# ─────────────────────────────────────────────────────────────────────────────

doc.add_heading("13. Conclusion et perspectives", level=1)

para("Le projet (RE)Sources Relationnelles a permis a notre equipe de mettre en pratique l'ensemble des competences du bloc INFCDAAL2 'Developper et tester les applications informatiques'. En partant du cahier des charges du Ministere des Solidarites et de la Sante, nous avons concu, developpe et teste une plateforme complete articulee autour d'une API REST securisee et de clients web et mobile.")

para("Points forts du projet :", bold=True)
bullet("Architecture MVC stricte avec Laravel, garantissant la maintenabilite et la testabilite du code")
bullet("Securite renforcee : chiffrement des donnees personnelles, anonymisation RGPD, headers de securite, rate limiting")
bullet("Frontend web complet avec accessibilite RGAA integree des la conception (skip links, ARIA, focus visible, contrastes)")
bullet("Suite de tests backend exhaustive couvrant l'authentification, les CRUD, les interactions, l'administration et la moderation")
bullet("Environnement Docker reproductible simplifiant le deploiement et la collaboration")
bullet("Organisation agile avec un workflow Git structure et des conventions de commits")

para("Perspectives d'evolution :", bold=True)
bullet("Enrichissement de l'application mobile avec des fonctionnalites avancees (notifications push, mode hors-ligne)")
bullet("Ajout de tests automatises frontend (Jest/Vitest) et de tests end-to-end (Cypress/Playwright)")
bullet("Implementation des fonctionnalites de messagerie et d'invitation entre participants")
bullet("Fonctionnalite d'export des statistiques (CSV, PDF)")
bullet("Support multilingue pour les contenus en langues etrangeres")
bullet("Mise en place d'un pipeline CI/CD complet (GitHub Actions) pour l'integration et le deploiement continus")

doc.add_paragraph()
para("Ce document technique, accompagne du cahier de tests et du cahier de recettes fournis en annexe, constitue la livraison documentaire complete du projet collaboratif (RE)Sources Relationnelles.",
     italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)


# =============================================================================
# FINALIZATION
# =============================================================================

add_page_numbers()
add_header_text()
doc.save(OUTPUT_PATH)
print(f"Document genere avec succes : {OUTPUT_PATH}")
print(f"Nombre de sections de niveau 1 : 13")
