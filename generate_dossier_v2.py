#!/usr/bin/env python3
"""
Generateur du Dossier Technique ALLEGE — (RE)Sources Relationnelles
CESI — Bloc 2 — Groupe 2 — INFCDAAL2
Version 2 : 17-19 pages (cible 15-20 pages consignes)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import datetime
import os

OUTPUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Dossier_Technique_v2.docx")
FONT_BODY = "Calibri"
PRIMARY_COLOR = RGBColor(0x1D, 0x3A, 0x6D)
GRAY_COLOR = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# =============================================================================
# HELPERS
# =============================================================================

def setup_styles():
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    for level, (size, color) in enumerate([
        (Pt(18), PRIMARY_COLOR),
        (Pt(14), PRIMARY_COLOR),
        (Pt(12), RGBColor(0x33, 0x33, 0x33)),
    ], start=1):
        h = doc.styles[f'Heading {level}']
        h.font.name = FONT_BODY
        h.font.size = size
        h.font.bold = True
        h.font.color.rgb = color
        h.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        h.paragraph_format.space_after = Pt(8)

def setup_margins():
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

def add_page_numbers():
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
        run2 = p.add_run()
        run2._r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
        run3 = p.add_run()
        run3._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

def add_header_text():
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("(RE)Sources Relationnelles — Dossier Technique")
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY_COLOR
        run.font.italic = True

def para(text, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, font_size=None, color=None, space_after=None):
    p = doc.add_paragraph()
    p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size:
        run.font.size = font_size
    if color:
        run.font.color.rgb = color
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="1D3A6D"/>'))
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F6FA"/>'))
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def visual_placeholder(vtype, section, content, tool, size, legend):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.text = ""
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFF9E6"/>'))
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("[ESPACE RESERVE — A COMPLETER]\n\n")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0x77, 0x00)
    details = (
        f"Type      : {vtype}\n"
        f"Section   : {section}\n"
        f"Contenu   : {content}\n"
        f"Outil     : {tool}\n"
        f"Taille    : {size}\n"
        f"Legende   : {legend}"
    )
    run2 = p.add_run(details)
    run2.font.size = Pt(9)
    run2.font.color.rgb = GRAY_COLOR
    run2.italic = True
    doc.add_paragraph()

def page_break():
    doc.add_page_break()


# =============================================================================
# GENERATION DU DOCUMENT ALLEGE
# =============================================================================

setup_styles()
setup_margins()

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ═══════════════════════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

para("(RE)Sources Relationnelles", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
     font_size=Pt(28), color=PRIMARY_COLOR)
doc.add_paragraph()
para("Dossier Technique", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
     font_size=Pt(18), color=GRAY_COLOR)
doc.add_paragraph()
para("Projet collaboratif — INFCDAAL2", alignment=WD_ALIGN_PARAGRAPH.CENTER,
     font_size=Pt(14), color=GRAY_COLOR)

for _ in range(3):
    doc.add_paragraph()

para("Plateforme de ressources pour l'amelioration\ndes relations entre citoyens",
     italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=Pt(12), color=GRAY_COLOR)

for _ in range(4):
    doc.add_paragraph()

add_table(
    ["Membre", "Perimetre"],
    [
        ["Chamil MUSAYEV", "Backend — API, BDD, Auth, Securite, RGPD"],
        ["Jamal", "Frontend web — Next.js, Charte graphique, RGAA"],
        ["Romain CONDOMITTI", "Back-office — Admin, Auth frontend, Integration"],
        ["Leon HEU", "Mobile — React Native, Expo"],
    ],
    col_widths=[6, 10]
)

para(f"CESI Ecole d'Ingenieurs — {datetime.date.today().strftime('%d/%m/%Y')}",
     alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=Pt(11), color=GRAY_COLOR)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# TABLE DES MATIERES
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("Table des matieres", level=1)
para("(Clic droit > Mettre a jour les champs > Mettre a jour toute la table)",
     italic=True, color=GRAY_COLOR, font_size=Pt(10))

p_toc = doc.add_paragraph()
run_toc = p_toc.add_run()
run_toc._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
run_toc2 = p_toc.add_run()
run_toc2._r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>'))
run_toc3 = p_toc.add_run()
run_toc3._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION  (allege : pas de contexte ministeriel, pas de perimetre)
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("1. Introduction", level=1)

para("Le projet (RE)Sources Relationnelles est une plateforme numerique proposant des ressources pour ameliorer la qualite des relations interpersonnelles des citoyens. Ce document technique presente les choix architecturaux, les solutions techniques retenues et la demarche de developpement mise en oeuvre par notre equipe.")

doc.add_heading("1.1 Acteurs du systeme", level=2)
para("Le systeme identifie cinq profils d'acteurs avec des niveaux de privileges distincts :")

add_table(
    ["Acteur", "Role", "Acces"],
    [
        ["Citoyen non connecte", "Consultation des ressources publiques", "Front-office (lecture seule)"],
        ["Citoyen connecte", "Creation, commentaires, favoris, progression", "Front-office (lecture/ecriture)"],
        ["Moderateur", "Validation des ressources et commentaires", "Front-office + moderation"],
        ["Administrateur", "Gestion du catalogue, statistiques, suspension", "Front-office + back-office"],
        ["Super-administrateur", "Creation de comptes privilegies", "Acces complet"],
    ],
    col_widths=[4, 6, 5]
)

doc.add_heading("1.2 Fonctionnalites implementees", level=2)
para("Le tableau suivant recapitule les fonctionnalites du CCTP et leur etat d'implementation :")

add_table(
    ["Domaine", "Fonctionnalite", "Statut"],
    [
        ["Acces ressources", "Listing pagine, filtres (categorie, type relation, type ressource), tri, detail", "Implemente"],
        ["Referentiel", "CRUD ressources en back-office, gestion categories, suspension", "Implemente"],
        ["Comptes", "Inscription citoyen, anonymisation RGPD, creation comptes privilegies", "Implemente"],
        ["Creation", "Creation/edition de ressources, validation par moderateur", "Implemente"],
        ["Echanges", "Commentaires hierarchiques, reponses, moderation", "Implemente"],
        ["Progression", "Favoris, ressources exploitees, mises de cote, tableau de bord", "Implemente"],
        ["Statistiques", "Tableau de bord admin avec filtres par periode et categorie", "Implemente"],
        ["Messagerie", "Invitation de participants, messagerie inter-utilisateurs", "Backlog"],
        ["Export stats", "Export des statistiques (CSV, PDF)", "Backlog"],
        ["Multilingue", "Traductions en langues etrangeres", "Backlog"],
    ],
    col_widths=[3, 8.5, 2.5]
)

para("Sur les fonctionnalites identifiees dans le CCTP, les parcours utilisateur principaux (consultation, creation, interaction, progression) ainsi que les circuits d'administration et de moderation sont pleinement implementes. Trois fonctionnalites avancees restent dans le backlog pour une version ulterieure.")

# ═══════════════════════════════════════════════════════════════════════════════
# 2. CHOIX TECHNOLOGIQUES  (allege : pas de tableau comparatif)
# ═══════════════════════════════════════════════════════════════════════════════

page_break()

doc.add_heading("2. Choix technologiques et justifications", level=1)

para("Chaque decision technique a ete prise en considerant les contraintes du CCTP, les competences de l'equipe et les bonnes pratiques du secteur.")

doc.add_heading("2.1 Backend : Laravel 11 (PHP 8.4)", level=2)
para("Le framework Laravel a ete retenu pour son architecture MVC native repondant directement a l'exigence du CCTP, son ORM Eloquent simplifiant la modelisation des relations entre entites, et Laravel Sanctum offrant une gestion de tokens API legere et securisee. Les FormRequest dedies permettent une validation declarative, et l'ecosysteme integre (migrations, seeders, testing PHPUnit) couvre l'ensemble du cycle de developpement.")

doc.add_heading("2.2 Frontend web : Next.js 16 (React 19, TypeScript)", level=2)
para("Le frontend repose sur Next.js 16 avec l'App Router pour un routage base sur le systeme de fichiers, React 19 pour ses ameliorations de performance et TypeScript pour le typage statique. Tailwind CSS v4 permet un prototypage rapide et un responsive design natif via les breakpoints (sm, md, lg, xl). Ce choix s'explique par l'expertise React de l'equipe et les besoins SEO des pages publiques (Server-Side Rendering).")

doc.add_heading("2.3 Base de donnees : PostgreSQL 16", level=2)
para("PostgreSQL a ete retenu pour son support natif du chiffrement (crucial pour l'implementation RGPD), ses types de donnees riches (text, jsonb), sa robustesse ACID et sa compatibilite native avec Laravel via le driver pdo_pgsql. L'image Docker legere postgres:16-alpine simplifie l'environnement de developpement.")

doc.add_heading("2.4 Application mobile : React Native (Expo)", level=2)
para("React Native avec Expo permet un developpement cross-platform (iOS et Android) a partir d'une base de code unique, avec un partage de logique metier avec le frontend web (meme ecosysteme React/TypeScript). Expo simplifie la configuration native et le hot reload accelere le cycle de developpement.")

doc.add_heading("2.5 Infrastructure : Docker Compose", level=2)
para("L'environnement est entierement containerise via Docker Compose avec quatre services : PostgreSQL 16 avec health checks et volume persistant, API Laravel (PHP 8.4) avec auto-migration, Adminer pour l'administration de la BDD, et Next.js (Node 22) avec hot reload. Cette approche garantit la reproductibilite entre les postes de developpement.")

add_table(
    ["Service", "Image", "Port", "Role"],
    [
        ["db", "postgres:16-alpine", "5432", "Base de donnees avec healthcheck et volume persistant"],
        ["api", "php:8.4-cli-alpine (custom)", "8000", "API Laravel avec migrations automatiques au demarrage"],
        ["db_admin", "adminer:4", "8080", "Interface web d'administration BDD"],
        ["web", "node:22-alpine", "3005", "Frontend Next.js avec hot reload et pnpm"],
    ],
    col_widths=[2, 5, 1.5, 7.5]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 3. ARCHITECTURE APPLICATIVE  (garde complet avec visuels)
# ═══════════════════════════════════════════════════════════════════════════════

page_break()

doc.add_heading("3. Architecture applicative", level=1)

doc.add_heading("3.1 Architecture globale du systeme", level=2)
para("L'architecture suit un modele client-serveur decouple : le backend expose une API REST JSON consommee independamment par le frontend web et l'application mobile. La base de donnees PostgreSQL est accessible uniquement par l'API, garantissant un point d'acces unique et securise.")

visual_placeholder(
    "Schema d'architecture globale",
    "3.1 Architecture globale",
    "Client Web (Next.js) et Client Mobile (React Native/Expo) communiquant via HTTP/JSON avec l'API REST (Laravel). L'API communique avec PostgreSQL. Docker Compose englobe API + DB + Adminer + Web.",
    "draw.io / Lucidchart",
    "Demi-page",
    "Figure 1 — Architecture globale du systeme"
)

doc.add_heading("3.2 Pattern MVC — Implementation Laravel", level=2)
para("Conformement au CCTP, l'architecture backend adopte strictement le pattern Modele-Vue-Controleur :")

add_table(
    ["Couche", "Responsabilite", "Implementation"],
    [
        ["Modele", "Donnees, relations, regles metier", "8 modeles Eloquent : User, Resource, Comment, Category, Favorite, Progression, RelationType, ResourceType"],
        ["Vue", "Presentation des donnees", "Reponses JSON (API-first). Les vues HTML sont gerees par le frontend Next.js."],
        ["Controleur", "Traitement des requetes", "10 controleurs organises par domaine : Auth, Resource, Comment, Favorite, Progression, Admin, Moderation, SuperAdmin, Category"],
    ],
    col_widths=[2, 4.5, 9.5]
)

doc.add_heading("3.3 Modele conceptuel de donnees", level=2)

visual_placeholder(
    "Modele Conceptuel de Donnees (MCD Merise)",
    "3.3 Modele conceptuel de donnees",
    "MCD Merise avec les entites : UTILISATEUR, RESSOURCE, COMMENTAIRE, CATEGORIE, TYPE_RELATION, TYPE_RESSOURCE, FAVORI, PROGRESSION. Relations avec cardinalites. Inclure les associations CREER (1,n), APPARTENIR (1,1), COMMENTER (0,n), AJOUTER_FAVORI (0,n), SUIVRE_PROGRESSION (0,n).",
    "Looping / JMerise / draw.io",
    "Pleine page",
    "Figure 2 — Modele Conceptuel de Donnees (MCD Merise)"
)

para("Le modele repose sur 8 tables principales :")

add_table(
    ["Table", "Role metier", "Colonnes principales"],
    [
        ["users", "Comptes utilisateurs, donnees chiffrees", "id, name (chiffre), email (chiffre), email_hash (SHA-256), password (hashe), role (enum), is_active"],
        ["resources", "Ressources avec workflow de validation", "id, title, content, user_id (FK), category_id (FK), relation_type_id (FK), resource_type_id (FK), status (enum), is_public"],
        ["comments", "Commentaires hierarchiques moderes", "id, content, user_id (FK), resource_id (FK), parent_id (FK self-ref), is_approved"],
        ["favorites", "Favoris utilisateur-ressource", "id, user_id (FK), resource_id (FK) — unique(user_id, resource_id)"],
        ["progressions", "Avancement utilisateur-ressource", "id, user_id (FK), resource_id (FK), status (enum) — unique(user_id, resource_id)"],
        ["categories", "Categories de ressources", "id, name, description"],
        ["relation_types", "Types de relations (couple, famille...)", "id, name"],
        ["resource_types", "Types de contenus (article, video...)", "id, name"],
    ],
    col_widths=[2.5, 3.5, 10]
)

para("Contraintes d'integrite : cles etrangeres CASCADE sur toutes les relations, unicite sur favorites et progressions (user_id, resource_id), index sur resources.status, resources.is_public et comments.is_approved, email_hash unique pour garantir l'unicite malgre le chiffrement.")

doc.add_heading("3.4 Diagramme de classes UML", level=2)

visual_placeholder(
    "Diagramme de classes UML",
    "3.4 Diagramme de classes",
    "8 entites avec attributs types, methodes et associations : User 1-->* Resource, User 1-->* Comment, User 1-->* Favorite, User 1-->* Progression, Resource *-->1 Category/RelationType/ResourceType, Resource 1-->* Comment, Comment 0..1-->* Comment (replies). Multiplicites et navigabilite.",
    "StarUML / draw.io / PlantUML",
    "Pleine page",
    "Figure 3 — Diagramme de classes UML"
)

para("Les relations Eloquent refletent ce diagramme : User hasMany Resource/Comment/Favorite/Progression ; Resource belongsTo User/Category/RelationType/ResourceType et hasMany Comment/Favorite/Progression ; Comment belongsTo User/Resource/parent et hasMany replies.")

# ═══════════════════════════════════════════════════════════════════════════════
# 4. DEVELOPPEMENT BACKEND  (allege : pas de table 24 endpoints)
# ═══════════════════════════════════════════════════════════════════════════════

page_break()

doc.add_heading("4. Developpement backend — API REST Laravel", level=1)

doc.add_heading("4.1 Organisation des routes API", level=2)
para("L'API REST est structuree autour de groupes de routes proteges par des middlewares progressifs. Toutes les reponses sont forcees en JSON.")

add_table(
    ["Groupe", "Middleware", "Routes principales"],
    [
        ["Public", "Aucun (+ throttle:auth sur register/login)", "GET /resources, GET /resources/{id}, GET /categories, GET /resources/{id}/comments, POST /register, POST /login"],
        ["Authentifie", "auth:sanctum", "POST /logout, GET /user, DELETE /user, POST /resources, PUT /resources/{id}, commentaires, favoris, progression"],
        ["Administration", "auth:sanctum + role:admin,super_admin", "GET /admin/statistics, PUT /admin/resources/{id}/suspend"],
        ["Moderation", "auth:sanctum + role:moderator,admin,super_admin", "PUT /moderation/resources/{id}/validate, commentaires approve/delete"],
        ["Super-admin", "auth:sanctum + role:super_admin", "POST /super-admin/users"],
    ],
    col_widths=[2.5, 5, 8.5]
)

doc.add_heading("4.2 Validation et sanitization des entrees", level=2)
para("Chaque endpoint d'ecriture utilise un FormRequest dedie assurant validation et sanitization via le trait SanitizesInput : suppression des balises HTML (strip_tags), normalisation des espaces, trim, normalisation des emails en minuscules, et hashage SHA-256 pour la recherche sans dechiffrement.")

add_table(
    ["FormRequest", "Champ", "Regles"],
    [
        ["RegisterRequest", "name / email / password", "required, string, min:2-255, email:rfc, confirmed, Password::defaults()"],
        ["StoreResourceRequest", "title / content / *_id", "required, string min:3-10, integer exists:table,id"],
        ["CommentRequest", "content", "required, string, min:2, max:2000"],
    ],
    col_widths=[3.5, 4, 8.5]
)

doc.add_heading("4.3 Flux d'authentification", level=2)
para("L'authentification repose sur Laravel Sanctum avec des tokens Bearer :")

visual_placeholder(
    "Diagramme de sequence — Authentification",
    "4.3 Flux d'authentification",
    "Client -> API (POST /login email+password) -> Recherche via email_hash -> Hash::check(password) -> Verification is_active -> createToken() -> Retour Bearer token + user. Puis : Client -> API (GET /user + Authorization: Bearer) -> Validation Sanctum -> Retour user.",
    "PlantUML / draw.io / Mermaid",
    "Demi-page",
    "Figure 4 — Diagramme de sequence du flux d'authentification"
)

para("L'email est stocke chiffre en base ; la recherche se fait via email_hash (SHA-256). Le mot de passe est verifie par Hash::check(), le flag is_active bloque les comptes desactives. Le token est transmis dans l'en-tete Authorization: Bearer et revoque a la deconnexion.")

doc.add_heading("4.4 Controle d'acces par roles", level=2)
para("Le controle d'acces combine un RoleMiddleware dynamique (verifie le role utilisateur contre la liste autorisee, retourne 403 en cas de refus) et une ResourcePolicy garantissant que seul l'auteur peut modifier sa ressource.")

add_table(
    ["Role", "Enum", "Droits"],
    [
        ["Citoyen", "citizen", "Consultation, creation (PENDING), commentaires, favoris, progression"],
        ["Moderateur", "moderator", "Validation ressources (PENDING -> PUBLISHED), approbation/suppression commentaires"],
        ["Administrateur", "admin", "Statistiques, suspension ressources (-> ARCHIVED), droits moderateur inclus"],
        ["Super-admin", "super_admin", "Creation comptes privilegies (moderateur, admin), tous les droits"],
    ],
    col_widths=[3, 2.5, 10.5]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 5. SECURITE ET RGPD  (allege : headers migres vers PPT)
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("5. Securite et conformite RGPD", level=1)

doc.add_heading("5.1 Mesures de securite implementees", level=2)

add_table(
    ["Mesure", "Implementation technique", "Fichier source"],
    [
        ["Hashage mot de passe", "Cast Laravel 'hashed' (bcrypt)", "app/Models/User.php"],
        ["Chiffrement donnees personnelles", "Crypt::encryptString() sur name et email", "app/Models/User.php (mutators)"],
        ["Lookup email securise", "email_hash = SHA-256(email normalise)", "User.php + migration"],
        ["Rate limiting", "throttle:auth (5 tentatives/min) sur register et login", "routes/api.php"],
        ["Headers de securite", "CSP, X-Frame-Options DENY, X-Content-Type-Options nosniff, Referrer-Policy", "Middleware/SecurityHeaders.php"],
        ["Sanitization", "strip_tags, trim, normalisation via SanitizesInput", "Requests/Concerns/SanitizesInput.php"],
    ],
    col_widths=[3.5, 6, 6.5]
)

doc.add_heading("5.2 Conformite RGPD — Chiffrement et anonymisation", level=2)

para("Chiffrement des donnees personnelles :", bold=True)
para("Les champs name et email du modele User sont chiffres a l'ecriture via Crypt::encryptString() et dechiffres a la lecture. Ce mecanisme est transparent grace aux accesseurs/mutateurs Eloquent. Un champ email_hash (SHA-256 de l'email normalise) permet l'authentification sans dechiffrer l'ensemble des emails.")

para("Droit a l'effacement — Anonymisation logique :", bold=True)
para("L'endpoint DELETE /api/user applique une anonymisation via User::anonymize() plutot qu'une suppression physique (qui briserait l'integrite referentielle) :")
bullet("Revocation de tous les tokens d'acces actifs")
bullet("Remplacement du nom par 'Deleted User' et de l'email par deleted-user-{id}-{uuid}@example.invalid")
bullet("Desactivation du compte (is_active = false), remise du role en citizen")
bullet("Effacement du mot de passe et de la date de verification email")

para("Ce choix preserve la coherence historique des contenus (ressources, commentaires) tout en rendant impossible l'identification de l'ancien utilisateur.")

# ═══════════════════════════════════════════════════════════════════════════════
# 6. DEVELOPPEMENT FRONTEND WEB  (allege : pas de charte graphique, pas d'admin)
# ═══════════════════════════════════════════════════════════════════════════════

page_break()

doc.add_heading("6. Developpement frontend web — Next.js", level=1)

doc.add_heading("6.1 Architecture de l'application web", level=2)
para("L'application est construite avec Next.js 16 (App Router) et organisee en deux espaces : un espace public pour les citoyens et un espace d'administration pour les roles privilegies.")

add_table(
    ["Route", "Page", "Description"],
    [
        ["/", "Accueil", "Hero, categories, fonctionnement, appel a l'action"],
        ["/ressources", "Listing ressources", "Recherche, filtres, tri, pagination (15/page)"],
        ["/ressources/[id]", "Detail ressource", "Contenu, commentaires hierarchiques, favoris"],
        ["/auth/connexion", "Connexion", "Formulaire avec gestion d'erreurs API"],
        ["/auth/inscription", "Inscription", "Validation champ par champ, confirmation mot de passe"],
        ["/presentation", "Presentation", "Mission de la plateforme"],
        ["/aide", "Aide / FAQ", "Centre d'aide"],
        ["/administration/*", "Back-office", "Stats, moderation, gestion ressources/utilisateurs/categories"],
    ],
    col_widths=[3.5, 3, 9.5]
)

doc.add_heading("6.2 Composants UI reutilisables", level=2)
para("Une bibliotheque de composants assure la coherence visuelle et l'accessibilite :")

add_table(
    ["Composant", "Variantes", "Accessibilite"],
    [
        ["Button", "primary, secondary, accent, outline — tailles sm/md/lg", "focus-visible:ring-2, disabled:opacity-50"],
        ["Input", "Avec label et erreur integres", "aria-invalid, aria-describedby, role='alert' sur erreurs"],
        ["Card", "Padding none/sm/md/lg, hover optionnel", "Structure semantique, transitions"],
        ["Badge", "primary, secondary, accent, gray, success, warning, error", "Texte descriptif (pas de dependance a la couleur seule)"],
    ],
    col_widths=[2.5, 6, 7.5]
)

doc.add_heading("6.3 Gestion de l'authentification cote client", level=2)
para("L'authentification est geree via un React Context (AuthContext) encapsulant l'etat utilisateur et les methodes login/register/logout. Le token Sanctum est stocke en localStorage et transmis automatiquement dans les en-tetes de chaque requete via un client HTTP centralise (api.ts).")

para("Flux : au chargement, le contexte verifie la presence d'un token et appelle GET /api/user pour rehydrater l'etat. En cas de token invalide, il est supprime. A la connexion, le token retourne est stocke et l'utilisateur redirige vers les ressources. A la deconnexion, le token est revoque cote serveur puis supprime du localStorage.")

# ═══════════════════════════════════════════════════════════════════════════════
# 7. ACCESSIBILITE RGAA
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("7. Accessibilite — Conformite RGAA", level=1)

para("Le RGAA constitue une exigence du CCTP. L'application integre les bonnes pratiques d'accessibilite des la conception :")

add_table(
    ["Critere RGAA", "Implementation", "Fichier"],
    [
        ["Navigation au clavier", "Lien d'evitement 'Aller au contenu principal' visible au focus", "(main)/layout.tsx"],
        ["Landmarks ARIA", "nav[aria-label], main[id=main-content], footer semantique", "Navbar.tsx, layout.tsx, Footer.tsx"],
        ["Langue de la page", "html[lang='fr'] sur le document racine", "app/layout.tsx"],
        ["Labels de formulaire", "Chaque input associe a un label via htmlFor + id", "components/ui/Input.tsx"],
        ["Etats d'erreur", "aria-invalid, aria-describedby, role='alert' aria-live='polite'", "Input.tsx, inscription/page.tsx"],
        ["Menu responsive", "aria-expanded, aria-controls liant bouton et menu mobile", "Navbar.tsx"],
        ["Elements decoratifs", "aria-hidden='true' sur les emojis", "page.tsx (accueil)"],
        ["Contraste", "Primary #1D3A6D sur blanc : ratio > 7:1 (AAA)", "globals.css"],
        ["Focus visible", "outline: 3px solid var(--primary) avec offset 2px", "globals.css"],
        ["Pagination", "aria-label='Aller a la page N', aria-current='page'", "ressources/page.tsx"],
        ["Filtres", "aria-label descriptif sur chaque select de filtrage", "ressources/page.tsx"],
    ],
    col_widths=[3, 6.5, 6.5]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 8. APPLICATION MOBILE  (allege : architecture seulement)
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("8. Application mobile — React Native / Expo", level=1)

para("L'application mobile est concue avec React Native et Expo pour un deploiement cross-platform iOS et Android. L'architecture consomme la meme API REST Laravel que le frontend web.")

visual_placeholder(
    "Schema d'architecture mobile",
    "8 Application mobile",
    "Ecrans : Accueil, Liste ressources, Detail, Profil, Connexion avec React Navigation (stack/tab). Communication API Laravel via HTTP. Stockage token via AsyncStorage.",
    "draw.io / Figma",
    "Demi-page",
    "Figure 5 — Architecture de l'application mobile"
)

add_table(
    ["Ecran", "Fonctionnalite", "Correspondance API"],
    [
        ["Accueil", "Acces rapide aux ressources", "GET /api/resources"],
        ["Liste ressources", "Listing avec filtres et recherche", "GET /api/resources?category=&sort="],
        ["Detail ressource", "Contenu, commentaires, favori", "GET /api/resources/{id}, GET /.../comments"],
        ["Connexion / Inscription", "Authentification", "POST /api/login, POST /api/register"],
        ["Profil", "Informations utilisateur", "GET /api/user, POST /api/logout"],
        ["Progression", "Favoris et exploitations", "GET /api/progression"],
    ],
    col_widths=[3.5, 4.5, 8]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 9. PLAN DE TESTS  (allege : 1 tableau fusionne au lieu de 4)
# ═══════════════════════════════════════════════════════════════════════════════

page_break()

doc.add_heading("9. Plan de tests", level=1)

doc.add_heading("9.1 Methodologie de test", level=2)

add_table(
    ["Niveau", "Type", "Outil", "Couverture"],
    [
        ["Backend — Feature", "Tests d'integration API", "PHPUnit + Laravel TestCase", "Auth, CRUD, interactions, admin, moderation"],
        ["Backend — Unitaire", "Tests modeles et services", "PHPUnit", "Modeles, enums, sanitization"],
        ["Frontend", "Tests manuels parcours", "Navigateur + DevTools", "Pages publiques, admin, formulaires"],
        ["Accessibilite", "Audit RGAA", "Lighthouse / axe DevTools", "Contraste, ARIA, navigation clavier"],
    ],
    col_widths=[3.5, 3.5, 4, 5]
)

para("La base de test utilise le trait RefreshDatabase de Laravel pour l'isolation entre tests. Les tokens Sanctum sont generes dans les tests via actingAs() ou createToken().")

doc.add_heading("9.2 Cahier de tests backend", level=2)
para("Les tests feature couvrent l'ensemble des cas metier critiques. Le tableau suivant synthetise les 27 tests implementes :")

add_table(
    ["Suite", "Tests", "Scenarios couverts"],
    [
        ["AuthEndpointsTest", "10", "Inscription (sanitization, chiffrement, token), login (email_hash), rate limiting (429 apres 5 echecs), logout (revocation token), profil, routes protegees (401), anonymisation RGPD, middleware roles (403), headers securite"],
        ["ResourceEndpointsTest", "6", "Listing pagine avec filtres et tri, detail avec relations, creation (PENDING, sanitize), update (auteur seul, 403 sinon), categories ordonnees"],
        ["InteractionEndpointsTest", "6", "Commentaires approuves avec arbre, creation/reponse (citoyen uniquement), favoris (ajout/retrait), progression (exploit, set-aside, dashboard)"],
        ["AdminEndpointsTest", "5", "Statistiques (filtres periode/categorie), suspension ressource, validation + moderation, creation compte privilegie (super-admin), protection routes admin"],
    ],
    col_widths=[3.5, 1, 11.5]
)

doc.add_heading("9.3 Recette et PV", level=2)
para("Le cahier de recettes et les proces-verbaux de tests sont formalises dans les documents annexes fournis :")
bullet("Cahier_de_Recettes.pdf : scenarios de recette fonctionnelle couvrant les parcours utilisateur")
bullet("Cahier_de_tests.pdf : matrice de tests avec cas nominaux, limites et erreurs")

# ═══════════════════════════════════════════════════════════════════════════════
# 10. ORGANISATION DU PROJET  (allege : equipe + deploiement, pas de git/commits)
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("10. Organisation du projet", level=1)

doc.add_heading("10.1 Equipe et repartition des roles", level=2)

add_table(
    ["Membre", "GitHub", "Role", "Perimetre technique"],
    [
        ["Chamil MUSAYEV", "@musayevchamil", "Backend Lead", "API Laravel, BDD PostgreSQL, Auth Sanctum, Securite, RGPD"],
        ["Jamal", "@jamal09102007", "Frontend Lead", "Next.js, charte graphique, accessibilite RGAA, integration API"],
        ["Romain CONDOMITTI", "@romain-condomitti", "Back-office", "Admin, Auth frontend, integration composants, moderation"],
        ["Leon HEU", "@leonheu", "Mobile Lead", "React Native, Expo, screens multi-plateformes"],
    ],
    col_widths=[3, 3.5, 2.5, 7]
)

doc.add_heading("10.2 Deploiement", level=2)
para("L'ensemble de l'environnement est containerise via Docker Compose. Au demarrage, le conteneur API execute automatiquement l'installation des dependances, la generation de la cle applicative, et les migrations de base de donnees. Le seeding genere 10 utilisateurs de demonstration (tous roles), 20 ressources et 15 commentaires pour les tests fonctionnels.")

# ═══════════════════════════════════════════════════════════════════════════════
# FIN
# ═══════════════════════════════════════════════════════════════════════════════

add_page_numbers()
add_header_text()
doc.save(OUTPUT_PATH)
print(f"Document genere : {OUTPUT_PATH}")
